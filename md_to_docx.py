import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Функция для перевода Markdown в DOCX с форматированием
def md_to_docx(md_path, docx_path):
    # Читаем файл Markdown
    with open(md_path, 'r', encoding='utf-8') as md_file:
        md_text = md_file.read()

    # Переводим Markdown в HTML
    html = markdown.markdown(md_text)

    # Парсим HTML для форматирования
    soup = BeautifulSoup(html, 'html.parser')

    # Создаем документ DOCX
    doc = Document()

    # Перебираем элементы HTML и добавляем их в DOCX
    for element in soup.recursiveChildGenerator():
        if element.name == 'h1':
            p = doc.add_heading(element.text.strip(), level=1)
        elif element.name == 'h2':
            p = doc.add_heading(element.text.strip(), level=2)
        elif element.name == 'h3':
            p = doc.add_heading(element.text.strip(), level=3)
        elif element.name == 'h4':
            p = doc.add_heading(element.text.strip(), level=4)
        elif element.name == 'p':
            p = doc.add_paragraph(element.text.strip())
        elif element.name == 'ul':
            for li in element.find_all('li'):
                doc.add_paragraph(li.text.strip(), style='List Bullet')
        elif element.name == 'ol':
            for li in element.find_all('li'):
                doc.add_paragraph(li.text.strip(), style='List Number')
        elif element.name == 'strong':
            p = doc.add_paragraph()
            run = p.add_run(element.text.strip())
            run.bold = True
        elif element.name == 'em':
            p = doc.add_paragraph()
            run = p.add_run(element.text.strip())
            run.italic = True
        elif element.name == 'blockquote':
            p = doc.add_paragraph(element.text.strip())
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p.paragraph_format.left_indent = Pt(20)
        elif element.name == 'code':
            p = doc.add_paragraph()
            run = p.add_run(element.text.strip())
            run.font.name = 'Courier New'
            run.font.size = Pt(10)

    # Сохраняем DOCX файл
    doc.save(docx_path)

# Использование скрипта
if __name__ == "__main__":
    md_to_docx('example.md', 'example.docx')
